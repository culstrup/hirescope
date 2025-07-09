"""
Document processing for resumes and attachments
Handles PDF, DOCX, TXT, and other formats
"""

import io
import re
from typing import List

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None  # type: ignore[assignment]

try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore[assignment]


class DocumentProcessor:
    """Processes various document formats to extract text"""

    def __init__(self) -> None:
        self.supported_formats = self._check_supported_formats()

    def _check_supported_formats(self) -> List[str]:
        """Check which formats are supported based on installed libraries"""
        formats = [".txt"]

        if PyPDF2:
            formats.append(".pdf")
        else:
            print("⚠️  PyPDF2 not installed - PDF support disabled")

        if Document:
            formats.append(".docx")
        else:
            print("⚠️  python-docx not installed - DOCX support disabled")

        return formats

    def extract_text(self, content: bytes, filename: str) -> str:
        """
        Extract text from document content

        Args:
            content: Raw file content
            filename: Original filename (for format detection)

        Returns:
            Extracted text or error message
        """
        if not content:
            return "[Empty file]"

        # Detect format from filename
        file_ext = filename.lower().split(".")[-1] if "." in filename else ""

        try:
            if file_ext == "pdf":
                return self._extract_pdf_text(content)
            elif file_ext == "docx":
                return self._extract_docx_text(content)
            elif file_ext in ["txt", "text"]:
                return self._extract_txt_text(content)
            elif file_ext == "doc":
                return "[Legacy DOC format - manual conversion needed]"
            else:
                return f"[Unsupported format: .{file_ext}]"

        except Exception as e:
            return f"[Extraction failed: {str(e)[:100]}]"

    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF"""
        if not PyPDF2:
            return "[PDF extraction unavailable - install PyPDF2]"

        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"

            # Check if we got meaningful text
            if len(text.strip()) > 50:
                return text.strip()
            else:
                # Likely an image-based PDF
                page_count = len(pdf_reader.pages)
                creator = (
                    pdf_reader.metadata.get("/Creator", "Unknown")
                    if pdf_reader.metadata
                    else "Unknown"
                )

                return (
                    f"[IMAGE-BASED PDF DETECTED]\n"
                    f"Creator: {creator}\n"
                    f"Pages: {page_count}\n"
                    f"This PDF contains images/scanned content that requires OCR for text extraction."
                )

        except Exception as e:
            return f"[PDF extraction error: {str(e)[:100]}]"

    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX"""
        if not Document:
            return "[DOCX extraction unavailable - install python-docx]"

        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)

            text_parts = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n".join(text_parts)

        except Exception as e:
            return f"[DOCX extraction error: {str(e)[:100]}]"

    def _extract_txt_text(self, content: bytes) -> str:
        """Extract text from TXT files"""
        try:
            # Try common encodings
            for encoding in ["utf-8", "latin-1", "cp1252", "iso-8859-1"]:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue

            # If all fail, use utf-8 with error handling
            return content.decode("utf-8", errors="ignore")

        except Exception as e:
            return f"[TXT extraction error: {str(e)[:100]}]"

    def get_text_preview(self, text: str, max_length: int = 500) -> str:
        """Get a clean preview of extracted text"""
        if not text or text.startswith("["):
            return text

        # Clean up whitespace
        text = re.sub(r"\s+", " ", text).strip()

        # Return preview
        if len(text) <= max_length:
            return text
        else:
            return text[:max_length] + "..."
